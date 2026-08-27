from __future__ import annotations

import io
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from markdown_model import SourceDocument, inline_text


FIXED_ZIP_TIME = (1980, 1, 1, 0, 0, 0)


def _east_asia_font(properties: Any, name: str) -> None:
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    fonts.set(qn("w:eastAsia"), name)


def _set_font(style: Any, western: str, east_asia: str, size: float) -> None:
    style.font.name = western
    style.font.size = Pt(size)
    _east_asia_font(style._element.get_or_add_rPr(), east_asia)


def _configure(document: Document, document_date: str, title: str) -> None:
    fixed_time = datetime.fromisoformat(document_date).replace(tzinfo=timezone.utc)
    props = document.core_properties
    props.title = title
    props.subject = "报警管理系统正式项目交付物"
    props.author = "报警管理系统项目组"
    props.last_modified_by = "报警管理系统项目组"
    props.keywords = "报警管理系统,项目交付"
    props.created = fixed_time
    props.modified = fixed_time
    props.revision = 1

    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    styles = document.styles
    _set_font(styles["Normal"], "Arial", "Microsoft YaHei", 10.5)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    styles["Normal"].paragraph_format.space_after = Pt(5)
    for name, size, color in (
        ("Title", 22, "17365D"),
        ("Heading 1", 18, "17365D"),
        ("Heading 2", 15, "1F4E78"),
        ("Heading 3", 12, "2F5597"),
        ("Heading 4", 11, "2F5597"),
    ):
        _set_font(styles[name], "Arial", "Microsoft YaHei", size)
        styles[name].font.color.rgb = RGBColor.from_string(color)
        styles[name].font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = 1
    run = footer.add_run(f"报警管理系统 · 文档日期 {document_date}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)
    _east_asia_font(run._element.get_or_add_rPr(), "Microsoft YaHei")
    footer.add_run(" · 第 ")
    _append_field(footer, "PAGE")
    footer.add_run(" / ")
    _append_field(footer, "NUMPAGES")
    footer.add_run(" 页")


def _append_field(paragraph: Any, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    code = OxmlElement("w:instrText")
    code.set(qn("xml:space"), "preserve")
    code.text = f" {instruction} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, code, separate, value, end))


def _segments(nodes: list[dict[str, Any]], bold: bool = False, italic: bool = False) -> list[tuple[str, bool, bool, bool]]:
    result: list[tuple[str, bool, bool, bool]] = []
    for node in nodes:
        kind = node["type"]
        if kind == "text":
            result.append((str(node.get("raw", "")), bold, italic, False))
        elif kind == "codespan":
            result.append((str(node.get("raw", "")), bold, italic, True))
        elif kind in {"softbreak", "linebreak"}:
            result.append(("\n", bold, italic, False))
        elif kind == "strong":
            result.extend(_segments(node.get("children", []), True, italic))
        elif kind == "emphasis":
            result.extend(_segments(node.get("children", []), bold, True))
        elif kind == "link":
            label = inline_text(node.get("children", []))
            target = str(node.get("attrs", {}).get("url", ""))
            value = label if not target or target == label else f"{label}（{target}）"
            result.append((value, bold, italic, False))
        else:
            raise ValueError(f"不支持的 DOCX 行内节点：{kind}")
    return result


def _write_inline(paragraph: Any, nodes: list[dict[str, Any]]) -> None:
    for text, bold, italic, code in _segments(nodes):
        chunks = text.split("\n")
        for index, chunk in enumerate(chunks):
            if index:
                paragraph.add_run().add_break(WD_BREAK.LINE)
            if not chunk:
                continue
            run = paragraph.add_run(chunk)
            run.bold = bold
            run.italic = italic
            run.font.name = "Consolas" if code else "Arial"
            _east_asia_font(run._element.get_or_add_rPr(), "Microsoft YaHei")
            if code:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(70, 70, 70)


def _table_rows(node: dict[str, Any]) -> list[list[dict[str, Any]]]:
    rows: list[list[dict[str, Any]]] = []
    for child in node.get("children", []):
        if child["type"] == "table_head":
            rows.append(child.get("children", []))
        elif child["type"] == "table_body":
            for row in child.get("children", []):
                rows.append(row.get("children", []))
    return rows


def _mark_repeat_header(row: Any) -> None:
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)


def _flatten_list_item(node: dict[str, Any]) -> list[dict[str, Any]]:
    inline_nodes: list[dict[str, Any]] = []
    for child in node.get("children", []):
        if child["type"] in {"block_text", "paragraph"}:
            if inline_nodes:
                inline_nodes.append({"type": "softbreak"})
            inline_nodes.extend(child.get("children", []))
        elif child["type"] == "list":
            for nested in child.get("children", []):
                if inline_nodes:
                    inline_nodes.append({"type": "softbreak"})
                inline_nodes.extend(_flatten_list_item(nested))
    return inline_nodes


def _render_blocks(document: Document, nodes: Iterable[dict[str, Any]], quote: bool = False) -> None:
    for node in nodes:
        kind = node["type"]
        if kind == "blank_line":
            continue
        if kind == "heading":
            level = int(node.get("attrs", {}).get("level", 1))
            paragraph = document.add_paragraph(style="Title" if level == 1 else f"Heading {min(level, 4)}")
            _write_inline(paragraph, node.get("children", []))
            continue
        if kind in {"paragraph", "block_text"}:
            paragraph = document.add_paragraph()
            if quote:
                paragraph.paragraph_format.left_indent = Cm(0.7)
                paragraph.paragraph_format.right_indent = Cm(0.4)
                paragraph.paragraph_format.space_before = Pt(3)
            _write_inline(paragraph, node.get("children", []))
            continue
        if kind == "block_code":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.right_indent = Cm(0.4)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(6)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F2F2F2")
            paragraph._p.get_or_add_pPr().append(shading)
            text = str(node.get("raw", "")).rstrip("\n")
            run = paragraph.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            _east_asia_font(run._element.get_or_add_rPr(), "Microsoft YaHei")
            continue
        if kind == "list":
            ordered = bool(node.get("attrs", {}).get("ordered", False))
            for item in node.get("children", []):
                paragraph = document.add_paragraph(style="List Number" if ordered else "List Bullet")
                _write_inline(paragraph, _flatten_list_item(item))
            continue
        if kind == "block_quote":
            _render_blocks(document, node.get("children", []), quote=True)
            continue
        if kind == "thematic_break":
            paragraph = document.add_paragraph()
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "B4C6E7")
            border.append(bottom)
            paragraph._p.get_or_add_pPr().append(border)
            continue
        if kind == "table":
            rows = _table_rows(node)
            if not rows:
                continue
            width = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            table.autofit = True
            for row_index, cells in enumerate(rows):
                row = table.rows[row_index]
                if row_index == 0:
                    _mark_repeat_header(row)
                for column_index, cell_node in enumerate(cells):
                    cell = row.cells[column_index]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    _write_inline(paragraph, cell_node.get("children", []))
                    if row_index == 0:
                        for run in paragraph.runs:
                            run.bold = True
                        shading = OxmlElement("w:shd")
                        shading.set(qn("w:fill"), "D9EAF7")
                        cell._tc.get_or_add_tcPr().append(shading)
            document.add_paragraph().paragraph_format.space_after = Pt(0)
            continue
        raise ValueError(f"不支持的 DOCX 块节点：{kind}")


def _canonicalize(package: bytes) -> bytes:
    source = io.BytesIO(package)
    output = io.BytesIO()
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(output, "w", compression=zipfile.ZIP_STORED) as target:
        for name in sorted(archive.namelist()):
            info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_STORED
            info.create_system = 3
            info.external_attr = 0o100644 << 16
            target.writestr(info, archive.read(name))
    return output.getvalue()


def render_docx(source: SourceDocument, document_date: str) -> bytes:
    document = Document()
    _configure(document, document_date, source.title)
    _render_blocks(document, source.ast)
    package = io.BytesIO()
    document.save(package)
    return _canonicalize(package.getvalue())
